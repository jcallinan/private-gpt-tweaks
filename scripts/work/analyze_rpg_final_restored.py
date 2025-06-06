import os
import re
import sys
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt

# === CONFIG ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TIMEOUT = 120
RETRIES = 2
DEBUG = True
FUZZY_THRESHOLD = 0.94
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_ROOT = "use_case_outputs"
LOG_ROOT = "logs"

# === UTILS ===
def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step)]

def run_ollama(model, prompt, retries=RETRIES):
    for attempt in range(retries):
        try:
            print(f"🤖 Running model: {model} (Attempt {attempt + 1})")
            result = subprocess.run(
                ["ollama", "run", model],
                input=prompt.encode(),
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                timeout=TIMEOUT
            )
            return result.stdout.decode().strip()
        except subprocess.TimeoutExpired:
            print(f"⏱️ Timeout on {model}, retrying...")
            time.sleep(5)
    return None

def normalize(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Tables Used": "## Entities Used / Tables Used",
        "## Entities Used": "## Entities Used / Tables Used"
    }
    for k, v in replacements.items():
        text = text.replace(k, v)
    return text

def fuzzy_dedupe(results):
    seen, deduped = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            deduped.append(r)
    return deduped

def is_valid(text):
    return all(kw in text for kw in ["## Identification", "## Description", "## Process Steps"])

def extract_title_id(text, program):
    match_title = re.search(r"^#\s*(.+)", text, re.MULTILINE)
    title = match_title.group(1).strip() if match_title else "Untitled"
    title_clean = re.sub(r'[^\w\s-]', '', title).replace(" ", "-")[:75]
    match_id = re.search(r"Use Case ID\*\*:\s*UC-[A-Z]+-[A-Z]*?(\d+)", text)
    uc_id = match_id.group(1) if match_id else "XXX"
    return f"UC-AP-{program}-{uc_id}-{title_clean}"

def format_docx(text, path):
    doc = Document()
    for line in text.splitlines():
        if line.strip().startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(line.strip().replace("## ", "").strip())
            run.bold = True
            run.font.size = Pt(12)
        else:
            doc.add_paragraph(line.strip())
    doc.save(path)

def build_prompt(template, chunk, program):
    return f"""{template}

You are analyzing **IBM RPG code from AP{program}**. Your job is to extract high-level business logic in a structured use case format.

[RPG CODE]
{chunk}
[END CODE]
"""

# === MAIN EXECUTION ===
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python script.py <path-to-rpg-file>")
        sys.exit(1)

    rpg_path = sys.argv[1]
    program = re.search(r"AP(\d+)", rpg_path.upper()).group(1)
    lines = read_lines(rpg_path)
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(lines)

    now = datetime.now()
    date_folder = os.path.join(OUTPUT_ROOT, f"usecases-{now.strftime('%Y-%m-%d')}")
    run_folder = f"ap{program}_{now.strftime('%Y%m%d_%H%M%S')}"
    outdir = os.path.join(date_folder, run_folder)
    docx_dir = os.path.join(outdir, "docx")
    logdir = os.path.join(LOG_ROOT, f"logs_{now.strftime('%Y%m%d_%H%M%S')}")
    os.makedirs(outdir, exist_ok=True)
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(logdir, exist_ok=True)

    all_results = []
    low_conf = []
    for i, chunk in enumerate(chunks):
        print(f"\n🔍 Chunk {i+1}/{len(chunks)}")
        prompt = build_prompt(template, chunk, program)
        result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)
        if result:
            result = normalize(result)
            if is_valid(result):
                fname = extract_title_id(result, program)
                md_path = os.path.join(outdir, fname + ".md")
                docx_path = os.path.join(docx_dir, fname + ".docx")
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)
                format_docx(result, docx_path)
                print(f"✅ Saved: {fname}")
                all_results.append(result)
            else:
                fname = f"low_conf_{i+1:02d}.md"
                path = os.path.join(outdir, fname)
                with open(path, "w", encoding="utf-8") as f:
                    f.write(result)
                low_conf.append(result)
                print(f"⚠️ Low confidence output saved.")
        else:
            fail_path = os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt")
            with open(fail_path, "w") as f:
                f.write("No output from model.")
            print(f"❌ Chunk {i+1} failed.")

    # Deduped summary
    deduped = fuzzy_dedupe(all_results)
    with open(os.path.join(outdir, "SUMMARY.md"), "w", encoding="utf-8") as f:
        for uc in deduped:
            f.write(uc + "\n\n" + "="*60 + "\n\n")

    if low_conf:
        with open(os.path.join(outdir, "LOW_CONFIDENCE.md"), "w", encoding="utf-8") as f:
            for lc in low_conf:
                f.write(lc + "\n\n" + "="*60 + "\n\n")
        format_docx("\n\n".join(low_conf), os.path.join(docx_dir, "LOW_CONFIDENCE.docx"))

    print(f"\n✅ Done! {len(deduped)} strong use cases, {len(low_conf)} low confidence.")
    print(f"📁 Output folder: {outdir}")
