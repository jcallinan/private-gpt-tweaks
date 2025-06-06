import os
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
import re
from docx import Document

PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 120
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
FUZZY_THRESHOLD = 0.94
DEBUG = False

RPG_FILES = [
    "AP160.rpg36.txt", "AP298.rpg36.txt", "AP105.rpg36.txt", "AP192.rpg36.txt",
    "AP290.rpg36.txt", "AP1099.rpg36.txt", "AP296.rpg36.txt", "AP991P.rpg36.txt",
    "AP315.rpg.txt", "AP316.rpg.txt", "AP760.rpg36.txt", "AP760P.rpg36.txt",
    "AP765.rpg36.txt", "AP765N.rpg36.txt", "AP780.rpg36.txt", "AP780P.rpg36.txt",
    "AP790.rpg36.txt", "AP910.rpgle.txt", "AP945.rpgle.txt", "AP945C.clp.txt",
    "AP3155.rpg.txt"
]

BASE_FOLDER = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC"
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_ROOT = "use_case_outputs"

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step) if i + size <= len(lines)]

def run_model(prompt, model):
    try:
        result = subprocess.run(["ollama", "run", model],
                                input=prompt.encode(),
                                stdout=subprocess.PIPE,
                                stderr=subprocess.PIPE,
                                timeout=TIMEOUT)
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize_headers(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, right in replacements.items():
        text = text.replace(wrong, right)
    return text

def fuzzy_match(a, b):
    return SequenceMatcher(None, a, b).ratio()

def extract_title_and_id(text, program):
    title = "Untitled"
    match = re.search(r"#\s*(.*?)\n", text)
    if match:
        title = match.group(1).strip()
    title = re.sub(r'[^a-zA-Z0-9\- ]', '', title)[:75].strip().replace(" ", "-")
    id_match = re.search(r"Use Case ID.*?UC-[^\n]+", text)
    use_case_id = id_match.group(0).split()[-1].replace("UC-", "") if id_match else "XXX"
    return f"UC-AP-{program}-{use_case_id}-{title or 'Untitled'}"

def build_prompt(template, chunk, program):
    return f"""{template}

You are analyzing IBM RPG code from program AP{program}. Extract only business logic in the structure below. Do not describe the code, do not summarize it. Fill out the sections and generate a flowchart if appropriate.

[RPG CODE]
{chunk}
[END CODE]
"""

def save_docx(content, path):
    doc = Document()
    for para in content.splitlines():
        p = doc.add_paragraph()
        if para.startswith("## "):
            p.add_run(para[3:]).bold = True
        elif para.startswith("**") and "**:" in para:
            parts = para.split("**:")
            if len(parts) == 2:
                p.add_run(parts[0].replace("**", "") + ":").bold = True
                p.add_run(" " + parts[1].strip())
            else:
                p.add_run(para)
        else:
            p.add_run(para)
    doc.save(path)

def process_file(filename, template, date_tag):
    full_path = os.path.join(BASE_FOLDER, filename)
    program = re.findall(r"AP(\d+)", filename.upper())[0]
    lines = read_lines(full_path)
    chunks = chunk_lines(lines)

    ts = datetime.now().strftime('%Y%m%d_%H%M%S')
    outdir = os.path.join(OUTPUT_ROOT, f"usecases-{date_tag}", f"ap{program}_{ts}")
    os.makedirs(outdir, exist_ok=True)
    raw_file = os.path.join(outdir, "RAW_OLLAMA_OUTPUT.md")
    log_file = os.path.join(outdir, "LOW_CONFIDENCE.docx")
    summary_md = os.path.join(outdir, "SUMMARY.md")

    good_cases, low_conf = [], []
    raw_log = []

    for i, chunk in enumerate(chunks):
        prompt = build_prompt(template, chunk, program)
        output = run_model(prompt, PRIMARY_MODEL) or run_model(prompt, FALLBACK_MODEL)
        if not output:
            continue
        output = normalize_headers(output)
        raw_log.append(f"# Chunk {i+1}\n\n{output}\n\n{'='*60}\n")

        if "Use Case ID" in output and "## Description" in output:
            title = extract_title_and_id(output, program)
            md_path = os.path.join(outdir, f"{title}.md")
            docx_path = os.path.join(outdir, f"{title}.docx")
            with open(md_path, "w", encoding="utf-8") as f:
                f.write(output)
            save_docx(output, docx_path)
            good_cases.append((title, output))
        else:
            low_conf.append(output)

    with open(raw_file, "w", encoding="utf-8") as f:
        f.writelines(raw_log)
    with open(summary_md, "w", encoding="utf-8") as f:
        for title, text in good_cases:
            f.write(f"# {title}\n\n{text}\n\n{'='*60}\n\n")
    if low_conf:
        doc = Document()
        for item in low_conf:
            doc.add_paragraph(item)
            doc.add_paragraph("=" * 50)
        doc.save(log_file)

    print(f"✅ {filename}: {len(good_cases)} good, {len(low_conf)} low confidence → {outdir}")

if __name__ == "__main__":
    date_tag = datetime.now().strftime('%Y-%m-%d')
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()

    for rpg in RPG_FILES:
        try:
            process_file(rpg, template, date_tag)
        except Exception as e:
            print(f"❌ Failed processing {rpg}: {e}")
