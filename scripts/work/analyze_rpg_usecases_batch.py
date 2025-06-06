import os
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
import re
from docx import Document

# === CONFIG ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 120
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"
FUZZY_THRESHOLD = 0.93
INCLUDE_AP200 = False

SOURCE_FILES = [
    "AP160.rpg36.txt", "AP298.rpg36.txt", "AP105.rpg36.txt", "AP192.rpg36.txt", "AP290.rpg36.txt",
    "AP1099.rpg36.txt", "AP296.rpg36.txt", "AP991P.rpg36.txt", "AP315.rpg.txt", "AP316.rpg.txt",
    "AP760.rpg36.txt", "AP760P.rpg36.txt", "AP765.rpg36.txt", "AP765N.rpg36.txt", "AP780.rpg36.txt",
    "AP780P.rpg36.txt", "AP790.rpg36.txt", "AP910.rpgle.txt", "AP945.rpgle.txt", "AP945C.clp.txt",
    "AP3155.rpg.txt"
]
SOURCE_BASE = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC"

# === UTILITIES ===

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def sliding_chunks(lines, size, overlap):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, context, program):
    context_block = f"\nContext code from AP200 (reference only):\n{context}\n" if context else ""
    return f"""{template}
You are analyzing IBM RPG source code for program **AP{program}**, which is part of the Accounts Payable system.

Extract a structured business use case using the template provided. Focus on:
- AP vouchers, check printing, vendor processing, GL validation, invoice details, 1099s.

Skip summaries. Format and label all sections exactly.

{context_block}
[RPG CODE]
{chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(["ollama", "run", model],
                                input=prompt.encode(),
                                stdout=subprocess.PIPE,
                                stderr=subprocess.PIPE,
                                timeout=TIMEOUT)
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, right in replacements.items():
        text = text.replace(wrong, right)
    return text

def fuzzy_dedupe(results):
    seen, final = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            final.append(r)
    return final

def extract_title(text, program):
    match_id = re.search(r"Use Case ID.*?(UC-[\w\-]+)", text, re.IGNORECASE)
    uc_id = match_id.group(1) if match_id else f"UC-AP-{program}-XXX"
    match_title = re.search(r"(?i)^#\s*(.*?)\n", text)
    title = match_title.group(1).strip() if match_title else "Untitled"
    safe_title = re.sub(r"[^\w\- ]", '', title).strip().replace(' ', '-')
    return f"{uc_id}-{safe_title[:75]}"

def save_docx(text, path):
    doc = Document()
    for line in text.splitlines():
        if line.startswith("###"):
            doc.add_paragraph(line.replace("###", "").strip(), style="Heading 2")
        elif line.startswith("**") and "**" in line[2:]:
            doc.add_paragraph(line, style="Normal")
        else:
            doc.add_paragraph(line)
    doc.save(path)

# === MAIN LOGIC ===

def process_file(file):
    program = re.findall(r"AP(\d+)", file.upper())[0]
    full_path = os.path.join(SOURCE_BASE, file)
    ap_lines = read_lines(full_path)
    ap200_lines = read_lines(os.path.join(SOURCE_BASE, "AP200.rpg36.txt"))[:40] if INCLUDE_AP200 else []
    chunks = sliding_chunks(ap_lines, CHUNK_SIZE, CHUNK_OVERLAP)
    template = open(TEMPLATE_FILE, "r").read()
    context = '\n'.join(ap200_lines)

    now = datetime.now()
    date_folder = f"usecases-{now.strftime('%Y-%m-%d')}"
    ts = now.strftime("%Y%m%d_%H%M%S")
    base_output = os.path.join(OUTPUT_BASE, date_folder, f"ap{program}_{ts}")
    os.makedirs(base_output, exist_ok=True)
    os.makedirs(os.path.join(base_output, "docx"), exist_ok=True)
    os.makedirs(os.path.join(base_output, "LOW_CONFIDENCE"), exist_ok=True)
    os.makedirs(os.path.join(base_output, "logs"), exist_ok=True)

    results = []
    for idx, chunk in enumerate(chunks):
        prompt = build_prompt(template, chunk, context, program)
        out = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)
        if not out:
            continue
        out = normalize(out)
        use_case_name = extract_title(out, program)
        base_path = os.path.join(base_output, f"{use_case_name}.md")

        if "voucher" in out.lower() or "check" in out.lower() or "invoice" in out.lower():
            results.append(out)
            with open(base_path, "w", encoding="utf-8") as f:
                f.write(out)
            save_docx(out, os.path.join(base_output, "docx", f"{use_case_name}.docx"))
            print(f"✅ {file} - Chunk {idx+1}: Saved {use_case_name}")
        else:
            lc_path = os.path.join(base_output, "LOW_CONFIDENCE", f"{use_case_name}.md")
            with open(lc_path, "w", encoding="utf-8") as f:
                f.write(out)
            print(f"⚠️ {file} - Chunk {idx+1}: Low confidence")

    results = fuzzy_dedupe(results)
    with open(os.path.join(base_output, "SUMMARY.md"), "w", encoding="utf-8") as f:
        for r in results:
            f.write(r + "\n\n" + "="*60 + "\n\n")
    print(f"\n✅ Finished {file}: {len(results)} good use cases saved\n")

# === RUN ALL ===

if __name__ == "__main__":
    for f in SOURCE_FILES:
        process_file(f)
