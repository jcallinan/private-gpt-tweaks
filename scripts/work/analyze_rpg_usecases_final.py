import os
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
import re
import sys
from docx import Document

PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TIMEOUT = 120
DEBUG = True
FUZZY_THRESHOLD = 0.94
MODULE_NAME = "AP"

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, program):
    return f"""{template}

You are analyzing IBM RPG code from AP{program}. Extract a structured business use case. 
Focus on logic related to: check printing, voucher processing, invoice validation, and payment logic.

Do NOT summarize. Do NOT label the language. Use this format:
# <Use Case Title>

## Identification
**Use Case ID**: UC-AP-{program}-XXX  
**Module Group**: Accounts Payable  
**Legacy Program Ref**: AP{program}.RPG  
**Version**: 1.0  
**Last Update**: <Date>  
**Last Update By**: <Name>  
**Created**: <Date>  
**Created By**: <Name>  
**Approved By**: <Name or ?>

... [all other headers here]

[RPG CODE]
{chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(["ollama", "run", model], input=prompt.encode(), stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=TIMEOUT)
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize_headers(text, program):
    text = re.sub(r"##\s+(Input Validation|Validation Rules)", "## Input Type Validation Checks", text)
    text = re.sub(r"##\s+(Entities Used|Tables Used)", "## Entities Used / Tables Used", text)
    text = re.sub(r"(\*\*Use Case ID\*\*: UC-AP-)(\d+)", rf"\1{program}-\2", text)
    return text

def extract_title(text, program):
    match_id = re.search(r"Use Case ID.*?UC-AP-[^\n]*", text, re.IGNORECASE)
    id_part = match_id.group(0).strip().split()[-1].replace("UC-", "") if match_id else "XXX"
    match_title = re.search(r"(?i)^#\s*(.*?)\n", text)
    title = match_title.group(1).strip() if match_title else "Print Check"
    title_clean = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-').strip('-')[:75]
    return f"UC-AP-{program}-{id_part}-{title_clean}"

def save_as_docx(content, path):
    doc = Document()
    for para in content.splitlines():
        if para.startswith("### "):
            doc.add_heading(para[4:].strip(), level=2)
        elif para.startswith("**") and para.endswith("**"):
            doc.add_paragraph(para, style="Heading3")
        else:
            doc.add_paragraph(para)
    doc.save(path)

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def is_valid_output(text):
    return "Use Case ID" in text and "## Description" in text

def format_narrative(text):
    sections = ["Identification", "Description", "Pre-Condition", "Post-Condition", "Entities Used / Tables Used", "Program Steps", "Tests Needed"]
    out = ["Use Case Template\n"]
    for sec in sections:
        match = re.search(f"##+\s+{re.escape(sec)}(.*?)((?=## )|\Z)", text, re.DOTALL | re.IGNORECASE)
        if match:
            out.append(f"### {sec}\n")
            out.append(match.group(1).strip() + "\n")
    return '\n'.join(out)

def process_chunk(i, chunk, template, outdir, logdir, rawfile, docxdir, program):
    prompt = build_prompt(template, chunk, program)
    result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)
    if result:
        result = normalize_headers(result, program)
        with open(rawfile, "a", encoding="utf-8") as rf:
            rf.write(f"\n\n# Chunk {i+1}\n\n{result}\n\n{'='*60}\n")
        if is_valid_output(result):
            filename = extract_title(result, program)
            with open(os.path.join(outdir, f"{filename}.md"), "w", encoding="utf-8") as f:
                f.write(result + "\n")
            save_as_docx(result, os.path.join(docxdir, f"{filename}.docx"))
            print(f"✅ Chunk {i+1}: {filename}")
            return result
        else:
            with open(os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt"), "w", encoding="utf-8") as f:
                f.write(result)
            print(f"⚠️ Chunk {i+1} saved to LOW_CONFIDENCE")
            return result
    else:
        print(f"❌ Chunk {i+1}: no output.")
    return None

# === MAIN ===
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❌ Please provide path to RPG source file.")
        sys.exit(1)

    path = sys.argv[1]
    lines = read_lines(path)
    program = re.findall(r"AP(\d+)", os.path.basename(path).upper())[0]
    chunks = chunk_lines(lines)
    template = open("use_case_template.md", "r", encoding="utf-8").read()

    now = datetime.now()
    date_dir = os.path.join("use_case_outputs", f"usecases-{now.strftime('%Y-%m-%d')}")
    ts = now.strftime("%Y%m%d_%H%M%S")
    outdir = os.path.join(date_dir, f"ap{program}_{ts}")
    logdir = os.path.join("logs", f"logs_{ts}")
    docxdir = os.path.join(outdir, "word_docs")
    rawfile = os.path.join(outdir, "RAW_OLLAMA_OUTPUT.md")

    os.makedirs(outdir, exist_ok=True)
    os.makedirs(logdir, exist_ok=True)
    os.makedirs(docxdir, exist_ok=True)

    results = []
    low_conf = []

    for i, chunk in enumerate(chunks):
        res = process_chunk(i, chunk, template, outdir, logdir, rawfile, docxdir, program)
        if res:
            if is_valid_output(res):
                results.append(res)
            else:
                low_conf.append(res)

    deduped = fuzzy_dedupe(results)
    with open(os.path.join(outdir, "SUMMARY.md"), "w", encoding="utf-8") as f:
        for uc in deduped:
            f.write(format_narrative(uc) + "\n\n" + "="*60 + "\n\n")
    if low_conf:
        doc = Document()
        for uc in low_conf:
            doc.add_paragraph(uc + "\n" + "="*50 + "\n")
        doc.save(os.path.join(outdir, "LOW_CONFIDENCE.docx"))

    print(f"\n✅ Done! {len(deduped)} strong use cases, {len(low_conf)} low confidence.\n📁 Output: {outdir}")
