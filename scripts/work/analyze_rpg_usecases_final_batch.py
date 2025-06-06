import os
import re
import sys
import time
from datetime import datetime
from difflib import SequenceMatcher
import subprocess
from docx import Document

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 120
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
FUZZY_THRESHOLD = 0.94
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"
INCLUDE_AP200 = False

# === ALL FILES TO PROCESS ===
ALL_RPG_FILES = [
    "AP160.rpg36.txt", "AP298.rpg36.txt", "AP105.rpg36.txt", "AP192.rpg36.txt", "AP290.rpg36.txt",
    "AP1099.rpg36.txt", "AP296.rpg36.txt", "AP991P.rpg36.txt", "AP315.rpg.txt", "AP316.rpg.txt",
    "AP760.rpg36.txt", "AP760P.rpg36.txt", "AP765.rpg36.txt", "AP765N.rpg36.txt", "AP780.rpg36.txt",
    "AP780P.rpg36.txt", "AP790.rpg36.txt", "AP910.rpgle.txt", "AP945.rpgle.txt", "AP945C.clp.txt",
    "AP3155.rpg.txt"
]
SOURCE_DIR = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC"

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i + size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, context=None, program="XXX"):
    context_block = f"\nReference context from AP200 (not primary logic):\n{context}\n" if context else ""
    return f"""{template}
You are analyzing IBM RPG code from program AP{program}.
Only extract meaningful business logic use cases — voucher handling, GL, invoice, check processing, 1099s.

Return ONLY structured output. Follow the format. No commentary.
{context_block}
[RPG CODE]
{chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize_headers(text, program):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    return re.sub(r"(Use Case ID\*\*: UC-AP)-(\d+)", rf"\1-{program}-\2", text)

def is_valid_output(text):
    return "Use Case ID" in text and "## Description" in text

def extract_title(text, program):
    match_id = re.search(r"Use Case ID.*?UC-AP.*?(\d+)", text)
    id_part = match_id.group(1).zfill(3) if match_id else "XXX"
    title_match = re.search(r"(?i)^#\s*(.*?)$", text, re.MULTILINE)
    title = title_match.group(1).strip() if title_match else "Untitled"
    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-')
    return f"UC-AP-{program}-{id_part}-{clean_title[:75]}"

def save_as_docx(content, path):
    doc = Document()
    for line in content.splitlines():
        if line.startswith("### "):
            doc.add_paragraph(line[4:], style='Heading2')
        elif line.startswith("**") and line.endswith("**"):
            doc.add_paragraph(line.strip("*"), style='Heading3')
        else:
            doc.add_paragraph(line)
    doc.save(path)

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def format_narrative(text):
    headers = ["Identification", "Description", "Pre-Condition", "Post-Condition",
               "Entities Used / Tables Used", "Program Steps", "Tests Needed"]
    result = ["Use Case Template\n"]
    for h in headers:
        m = re.search(f"##+\s+{re.escape(h)}(.*?)(?=\n## |$)", text, re.DOTALL)
        if m:
            result.append(f"### {h}\n{m.group(1).strip()}\n")
    return "\n".join(result)

def process_chunks(file_path, program):
    now = datetime.now()
    date_folder = os.path.join(OUTPUT_BASE, f"usecases-{now.strftime('%Y-%m-%d')}")
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(date_folder, f"{program}_{timestamp}")
    log_dir = os.path.join("logs", f"logs_{program}_{timestamp}")
    low_conf_dir = os.path.join(output_dir, "LOW_CONFIDENCE")
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(log_dir, exist_ok=True)
    os.makedirs(low_conf_dir, exist_ok=True)

    raw_out = os.path.join(output_dir, "RAW_OLLAMA_OUTPUT.md")
    summary_md = os.path.join(output_dir, "SUMMARY.md")
    low_conf_docx = os.path.join(output_dir, "LOW_CONFIDENCE.docx")
    failed_txt = os.path.join(output_dir, "FAILED_CHUNKS.txt")

    all_lines = read_lines(file_path)
    ap200_lines = read_lines(os.path.join(SOURCE_DIR, "AP200.rpg36.txt")) if INCLUDE_AP200 else []
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(all_lines)

    all_results = []
    low_conf_results = []
    failed = []
    raw_log = []

    for i, chunk in enumerate(chunks):
        prompt = build_prompt(template, chunk, "\n".join(ap200_lines), program)
        result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)
        if result:
            result = normalize_headers(result, program)
            raw_log.append(f"\n\n# Chunk {i + 1}\n{result}\n{'=' * 50}\n")
            if is_valid_output(result):
                filename = extract_title(result, program)
                md_path = os.path.join(output_dir, f"{filename}.md")
                docx_path = os.path.join(output_dir, f"{filename}.docx")
                open(md_path, "w", encoding="utf-8").write(result)
                save_as_docx(result, docx_path)
                all_results.append(result)
            else:
                lowfile = os.path.join(low_conf_dir, f"low_conf_chunk_{i + 1}.md")
                open(lowfile, "w", encoding="utf-8").write(result)
                low_conf_results.append(result)
        else:
            failed.append(i + 1)

    with open(raw_out, "w", encoding="utf-8") as f:
        f.writelines(raw_log)

    with open(summary_md, "w", encoding="utf-8") as f:
        for uc in fuzzy_dedupe(all_results):
            f.write(format_narrative(uc) + "\n\n" + "=" * 60 + "\n\n")

    with open(failed_txt, "w") as f:
        for idx in failed:
            f.write(f"Chunk {idx} failed\n")

    if low_conf_results:
        doc = Document()
        for r in low_conf_results:
            doc.add_paragraph("========================================", style='Heading2')
            for line in r.splitlines():
                doc.add_paragraph(line)
        doc.save(low_conf_docx)

    print(f"✅ Done: {len(all_results)} strong use cases, {len(low_conf_results)} low confidence, {len(failed)} failed.")

def run_all():
    for fname in ALL_RPG_FILES:
        full_path = os.path.join(SOURCE_DIR, fname)
        program_match = re.search(r"AP(\d+)", fname.upper())
        if not program_match:
            print(f"Skipping {fname}: No AP### match.")
            continue
        process_chunks(full_path, program_match.group(1))

if __name__ == "__main__":
    if "--all" in sys.argv:
        run_all()
    elif len(sys.argv) >= 2:
        file_arg = sys.argv[1]
        program_match = re.search(r"AP(\d+)", file_arg.upper())
        if not program_match:
            print("❌ Invalid file name. Must include AP###.")
            sys.exit(1)
        process_chunks(file_arg, program_match.group(1))
    else:
        print("Usage: python analyze_rpg_usecases_final_batch.py [filename] or --all")
