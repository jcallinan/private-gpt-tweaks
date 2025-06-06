import os
import re
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TIMEOUT = 120
FUZZY_THRESHOLD = 0.94
DEBUG = True

PRIMARY_FILE = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC\AP160.rpg36.txt"
TEMPLATE_FILE = "use_case_template.md"
CONTEXT_FILE = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC\AP200.rpg36.txt"
INCLUDE_CONTEXT = False


# === FUNCTIONS ===
def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size, overlap):
    step = size - overlap
    return ['\n'.join(lines[i:i + size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, context=None):
    context_block = f"\nContext code from AP200 (reference only):\n{context}\n" if context else ""
    return f"""{template}

You are analyzing **IBM RPG code from AP160** (Accounts Payable system). Your job is to extract only the business logic in a structured use case format. Focus on:
- AP vouchers, vendor validation, invoice rules, GL, payments, and 1099 logic.
- Skip implementation detail and low-level RPG structure.

DO NOT return summaries or commentary. Follow the format strictly.
{context_block}
[RPG CODE]
{chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize_headers(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    return text

def is_valid_output(text):
    return all(section in text for section in [
        "Use Case ID", "## Description", "## Process Steps"
    ])

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def extract_title_and_id(text):
    id_match = re.search(r"Use Case ID.*?UC-AP-160-([^\s]+)", text)
    title_match = re.search(r"(?i)^#\s*(.*?)\n", text)
    use_case_id = id_match.group(1) if id_match else "XXX"
    title = title_match.group(1).strip() if title_match else "Untitled"
    title_clean = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-').strip('-')[:75]
    return use_case_id, title_clean

def save_docx(content, docx_path):
    doc = Document()
    for line in content.splitlines():
        if line.strip().startswith("## "):
            doc.add_paragraph(line.replace("## ", "").strip(), style='Heading2')
        elif line.strip().startswith("**") and "**:" in line:
            bold_key, val = line.split("**:", 1)
            p = doc.add_paragraph()
            p.add_run(bold_key.strip(" *") + ": ").bold = True
            p.add_run(val.strip())
        else:
            doc.add_paragraph(line.strip())
    doc.save(docx_path)

def format_narrative(text):
    sections = [
        "Identification", "Description", "Pre-Condition", "Post-Condition",
        "Entities Used / Tables Used", "Program Steps", "Tests Needed"
    ]
    out = ["Use Case Template\n"]
    for sec in sections:
        match = re.search(f"##+\s+{re.escape(sec)}(.*?)((?=## )|\Z)", text, re.DOTALL | re.IGNORECASE)
        if match:
            out.append(f"### {sec}\n")
            out.append(match.group(1).strip() + "\n")
    return '\n'.join(out)

def process_chunk(i, chunk, template, context, rawfile, outdir, docx_dir, logdir):
    prompt = build_prompt(template, chunk, context)
    result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)
    if not result:
        with open(os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt"), "w") as f:
            f.write("Model failed to return output.")
        print(f"❌ Chunk {i+1}: no output.")
        return None

    result = normalize_headers(result)

    with open(rawfile, "a", encoding="utf-8") as f:
        f.write(f"\n\n# Chunk {i+1}\n\n{result}\n\n{'='*60}\n")

    if is_valid_output(result):
        uc_id, uc_title = extract_title_and_id(result)
        filename_base = f"UC-AP-160-{uc_id}-{uc_title}"
        md_path = os.path.join(outdir, f"{filename_base}.md")
        docx_path = os.path.join(docx_dir, f"{filename_base}.docx")
        with open(md_path, "w", encoding="utf-8") as f:
            f.write(result + "\n")
        save_docx(result, docx_path)
        print(f"✅ Chunk {i+1}: Saved as {filename_base}")
        return result
    else:
        with open(os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt"), "w", encoding="utf-8") as f:
            f.write(result)
        print(f"❌ Chunk {i+1}: failed format check.")
        return None


# === MAIN RUN ===
if __name__ == "__main__":
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    date_folder = datetime.now().strftime("%Y-%m-%d")
    program = "ap160"
    base_folder = os.path.join("use_case_outputs", f"usecases-{date_folder}", f"{program}_{ts}")
    docx_dir = os.path.join(base_folder, "word_docs")
    log_dir = os.path.join("logs", f"logs_{ts}")
    rawfile = os.path.join(base_folder, "RAW_OLLAMA_OUTPUT.md")
    summary_path = os.path.join(base_folder, "SUMMARY.md")

    os.makedirs(base_folder, exist_ok=True)
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(log_dir, exist_ok=True)

    lines = read_lines(PRIMARY_FILE)
    context = "\n".join(read_lines(CONTEXT_FILE)[:40]) if INCLUDE_CONTEXT else ""
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(lines, CHUNK_SIZE, CHUNK_OVERLAP)

    all_results = []
    for i, chunk in enumerate(chunks):
        res = process_chunk(i, chunk, template, context, rawfile, base_folder, docx_dir, log_dir)
        if res:
            all_results.append(res)

    final = fuzzy_dedupe(all_results)
    with open(summary_path, "w", encoding="utf-8") as f:
        for uc in final:
            f.write(format_narrative(uc))
            f.write("\n\n" + "="*60 + "\n\n")

    print(f"\n✅ {len(final)} unique use cases saved to {base_folder}")
