import os
import re
import sys
import time
from datetime import datetime
from difflib import SequenceMatcher
import subprocess
from docx import Document

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 60  # Reduced timeout for faster processing
CHUNK_SIZE = 90  # Original size
CHUNK_OVERLAP = 30  # Original overlap
FUZZY_THRESHOLD = 0.94
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"
INCLUDE_AP200 = False  # Simplified - no extra context
MIN_ACCEPTANCE_SCORE = 0.5  # Threshold for accepting use cases even if not perfect

# === ALL FILES TO PROCESS ===
ALL_RPG_FILES = [
    "AP160.rpg36.txt", "AP298.rpg36.txt", "AP105.rpg36.txt", "AP192.rpg36.txt", "AP290.rpg36.txt",
    "AP1099.rpg36.txt", "AP296.rpg36.txt", "AP991P.rpg36.txt", "AP315.rpg.txt", "AP316.rpg.txt",
    "AP760.rpg36.txt", "AP760P.rpg36.txt", "AP765.rpg36.txt", "AP765N.rpg36.txt", "AP780.rpg36.txt",
    "AP780P.rpg36.txt", "AP790.rpg36.txt", "AP910.rpgle.txt", "AP945.rpgle.txt", "AP945C.clp.txt",
    "AP3155.rpg.txt"
]
SOURCE_DIR = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC"

def read_lines(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return [line.rstrip("\n") for line in f.readlines()]
    except UnicodeDecodeError:
        # Fallback to another encoding if utf-8 fails
        with open(path, "r", encoding="latin-1") as f:
            return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i + size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, context=None, program="XXX"):
    context_block = f"\nReference context from AP200 (not primary logic):\n{context}\n" if context else ""
    
    # Updated prompt to match the desired output format
    return f"""You are analyzing IBM RPG code to extract business logic use cases.
Review this code from program AP{program} and extract meaningful business use cases.

Your output MUST follow this EXACT template:

**[Clear Descriptive Title]**

**Identification**

**Use Case ID:** UC-AP-{program}-XX (where XX is a sequential number starting at 01)

**Module Group:** Accounts Payable

**Legacy Program Ref:** AP{program}.RPG36

**Version**: 1.0

**Last Update:** {datetime.now().strftime('%m.%d.%Y')}

**Last Update By:** System Generated

**Created:** {datetime.now().strftime('%m.%d.%Y')}

**Created By:** System Generated

**Approved By**: ?

**Description:**
[Write a clear description of the use case, what it does, and its purpose]

**Pre-Condition:**
[List any prerequisites or conditions that must be met before this process can run]

**Post-Condition:**
[List the outcomes and state changes after the process completes]

**Entities Used / Tables Used:**
[List key database tables and entities used]

**Process Steps:**
[List the key process steps in clear business terms]

**Tests Needed:**
[List key test scenarios]

**Business Rules:**
[List important business rules and validations]

{context_block}
[RPG CODE]
{chunk}
[END CODE]

Be concise but thorough. Focus on meaningful business logic related to vouchers, GL, invoices, payment processing, etc.
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        print(f"Timeout using {model}")
        return None
    except Exception as e:
        print(f"Error running ollama: {e}")
        return None

def normalize_headers(text, program):
    """
    Normalize the headers in the use case to match the required template format
    """
    # Basic cleanups
    text = re.sub(r"# (.+)", r"**\1**", text)  # Convert main title from markdown to bold
    
    # Ensure all section headers use the correct format
    sections = [
        "Identification", "Description", "Pre-Condition", "Post-Condition",
        "Entities Used / Tables Used", "Process Steps", "Tests Needed", "Business Rules"
    ]
    
    for section in sections:
        # Replace any variant of section headers with the correct format
        text = re.sub(rf"##\s+{re.escape(section)}s?", f"**{section}:**", text)
        text = re.sub(rf"###\s+{re.escape(section)}s?", f"**{section}:**", text)
        
        # If the section doesn't exist at all, add it
        if f"**{section}:**" not in text and section != "Business Rules":
            text += f"\n\n**{section}:**\n- TBD"
    
    # Fix Use Case ID format
    text = re.sub(r"Use Case ID\*\*: UC-AP-(\d+)-(\d{3})", r"Use Case ID:** UC-AP-\1-\2", text)
    
    # Make sure use case ID is properly formatted with 2-digit sequence numbers
    text = re.sub(r"Use Case ID\*\*: UC-AP-(\d+)-(\d{3})", 
                  lambda m: f"Use Case ID:** UC-AP-{m.group(1)}-{int(m.group(2)):02d}", text)
    
    # Ensure Module Group is present
    if "**Module Group:**" not in text:
        text = re.sub(r"(\*\*Use Case ID:\*\* [^\n]+)", 
                     r"\1\n\n**Module Group:** Accounts Payable", text)
    
    # Add other required fields if missing
    required_fields = [
        f"**Legacy Program Ref:** AP{program}.RPG36",
        "**Version**: 1.0",
        f"**Last Update:** {datetime.now().strftime('%m.%d.%Y')}",
        "**Last Update By:** System Generated",
        f"**Created:** {datetime.now().strftime('%m.%d.%Y')}",
        "**Created By:** System Generated",
        "**Approved By**: ?"
    ]
    
    for field in required_fields:
        if field.split(":")[0] not in text:
            # Insert after Identification section
            text = re.sub(r"(\*\*Module Group:[^\n]+)", 
                          lambda m: f"{m.group(1)}\n\n{field}", text)
    
    return text

def calculate_acceptance_score(text):
    """Calculate a score for how complete/valid the use case is based on the new template format"""
    score = 0
    required_sections = [
        "**Use Case ID:**", 
        "**Module Group:**", 
        "**Legacy Program Ref:**",
        "**Description:**", 
        "**Pre-Condition:**", 
        "**Post-Condition:**",
        "**Entities Used / Tables Used:**", 
        "**Process Steps:**"
    ]
    
    # Check if there's a title (must be in bold at beginning)
    if re.search(r"^\*\*[^*]+\*\*", text.strip()):
        score += 0.1
    
    # Check sections
    for section in required_sections:
        if section in text:
            score += 0.1
            
            # Check content in sections (look for text after section heading)
            pattern = f"{re.escape(section)}(.*?)(?=\*\*[^:]+:\*\*|$)"
            match = re.search(pattern, text, re.DOTALL)
            if match and len(match.group(1).strip()) > 5:  # Some minimal content
                score += 0.05
    
    return score

def is_valid_output(text):
    """Basic validation to check if this looks like a use case"""
    score = calculate_acceptance_score(text)
    return score >= MIN_ACCEPTANCE_SCORE

def extract_title(text, program, use_case_counter):
    """
    Extract title for filename based on the template format
    """
    # Get the title from the content - it should be the first bold text
    title_match = re.search(r"^\*\*([^*]+)\*\*", text.strip())
    
    if title_match:
        title = title_match.group(1).strip()
    else:
        # If no title found, use a generic one
        title = "Use Case"
    
    # Ensure use case counter is formatted with 2 digits
    id_part = f"{use_case_counter:02d}"
    
    # Clean the title for use in filename
    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-')
    clean_title = re.sub(r'-+', '-', clean_title)
    
    # Create the filename
    return f"UC-AP-{program}-{id_part}-{clean_title[:40]}"

def save_as_docx(content, path):
    try:
        doc = Document()
        
        # Process content as paragraphs with proper formatting
        paragraphs = content.split('\n\n')
        
        for para in paragraphs:
            lines = para.strip().split('\n')
            
            # If this is just a single line
            if len(lines) == 1:
                line = lines[0].strip()
                
                # Title (in bold)
                if line.startswith('**') and line.endswith('**') and ':' not in line:
                    # Remove the asterisks and add as heading
                    clean_title = line.strip('*').strip()
                    doc.add_heading(clean_title, level=1)
                # Section headers
                elif line.startswith('**') and ':**' in line:
                    # Remove the asterisks but keep the bold formatting
                    clean_header = line.replace('**', '').strip()
                    p = doc.add_paragraph()
                    p.add_run(clean_header).bold = True
                # Regular paragraph
                else:
                    doc.add_paragraph(line)
            # Multi-line paragraph
            else:
                # Check if first line is a section header
                first_line = lines[0].strip()
                if first_line.startswith('**') and ':**' in first_line:
                    # Remove the asterisks but keep the bold formatting
                    clean_header = first_line.replace('**', '').strip()
                    p = doc.add_paragraph()
                    p.add_run(clean_header).bold = True
                    
                    # Add remaining lines as separate paragraphs
                    for line in lines[1:]:
                        if line.strip():
                            if line.strip().startswith('-'):
                                doc.add_paragraph(line.strip(), style='List Bullet')
                            else:
                                doc.add_paragraph(line.strip())
                else:
                    # Just a regular multi-line paragraph
                    p = doc.add_paragraph()
                    for line in lines:
                        p.add_run(line + '\n')
        
        doc.save(path)
        return True
    except Exception as e:
        print(f"Error saving Word document: {e}")
        return False

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def format_narrative(text):
    """
    Format the use case narrative for inclusion in the summary document
    """
    # Keep the full text as is - no need to reformat since we're already in the desired format
    return text

def process_chunks(file_path, program):
    now = datetime.now()
    date_folder = os.path.join(OUTPUT_BASE, f"usecases-{now.strftime('%Y-%m-%d')}")
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(date_folder, f"{program}_{timestamp}")
    log_dir = os.path.join("logs", f"logs_{program}_{timestamp}")
    
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(log_dir, exist_ok=True)

    raw_out = os.path.join(output_dir, "RAW_OLLAMA_OUTPUT.md")
    summary_md = os.path.join(output_dir, "SUMMARY.md")
    summary_docx = os.path.join(output_dir, "SUMMARY.docx")
    failed_txt = os.path.join(output_dir, "FAILED_CHUNKS.txt")

    print(f"Processing {file_path} for program AP{program}...")
    
    all_lines = read_lines(file_path)
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(all_lines)
    
    print(f"Generated {len(chunks)} chunks to process")
    
    all_results = []
    failed = []
    raw_log = []
    use_case_counter = 1  # To ensure sequential numbering starts at 01

    # Process each chunk
    for i, chunk in enumerate(chunks):
        print(f"Processing chunk {i+1}/{len(chunks)}")
        
        # Build and run prompt
        prompt = build_prompt(template, chunk, None, program)
        
        # Try primary model first
        print(f"Trying with {PRIMARY_MODEL}")
        result = run_ollama(PRIMARY_MODEL, prompt)
        
        # Fall back to secondary model if needed
        if not result:
            print(f"Falling back to {FALLBACK_MODEL}")
            result = run_ollama(FALLBACK_MODEL, prompt)
        
        # Process the result
        if result:
            # Normalize headers to match template
            result = normalize_headers(result, program)
            raw_log.append(f"\n\n# Chunk {i + 1}\n{result}\n{'=' * 50}\n")
            
            # Check if it looks valid enough
            score = calculate_acceptance_score(result)
            
            if is_valid_output(result):
                # Make sure Use Case ID is in the right format with 2-digit sequence numbers
                
                # Find the ID pattern and replace it with the correctly formatted version
                id_pattern = f"Use Case ID**: UC-AP-{program}-"
                id_pattern_index = result.find(id_pattern)

                if id_pattern_index >= 0:
                    # Find where the ID ends (look for next whitespace or line break)
                    end_pos = result.find("\n", id_pattern_index)
                    if end_pos < 0:
                        end_pos = len(result)
                    
                    # Create the new ID string
                    new_id = f"{id_pattern}{use_case_counter:02d}"
                    
                    # Replace just that portion
                    result = result[:id_pattern_index] + new_id + result[end_pos:]
                
                # Generate filename with proper format
                filename = extract_title(result, program, use_case_counter)
                
                md_path = os.path.join(output_dir, f"{filename}.md")
                docx_path = os.path.join(output_dir, f"{filename}.docx")
                
                # Save as markdown
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)
                
                # Save as docx
                if save_as_docx(result, docx_path):
                    print(f"✅ Created use case: {filename}")
                    all_results.append(result)
                    use_case_counter += 1
                else:
                    print(f"⚠️ Created use case markdown but Word doc failed: {filename}")
                    all_results.append(result)  # Still add it to results
                    use_case_counter += 1
            else:
                # For low confidence results, still try to create a use case
                # Use a clearly marked filename
                base_filename = f"UC-AP-{program}-{use_case_counter:02d}-Low-Confidence"
                
                # Extract any title if possible
                title_match = re.search(r"^\*\*([^*]+)\*\*", result.strip())
                if title_match:
                    title = title_match.group(1).strip()
                    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-')
                    base_filename = f"UC-AP-{program}-{use_case_counter:02d}-{clean_title[:40]}"
                
                md_path = os.path.join(output_dir, f"{base_filename}.md")
                docx_path = os.path.join(output_dir, f"{base_filename}.docx")
                
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)
                    
                save_as_docx(result, docx_path)
                print(f"⚠️ Created low confidence use case: {base_filename}")
                use_case_counter += 1
        else:
            print(f"❌ Failed to process chunk {i+1}")
            failed.append(i + 1)
    
    # Write aggregated outputs
    print(f"Writing summary files...")
    
    # Raw output log
    with open(raw_out, "w", encoding="utf-8") as f:
        f.writelines(raw_log)
    
    # Deduped summary
    unique_results = fuzzy_dedupe(all_results)
    with open(summary_md, "w", encoding="utf-8") as f:
        for uc in unique_results:
            f.write(uc + "\n\n" + "=" * 60 + "\n\n")
    
    # Summary as Word document
    try:
        if unique_results:
            doc = Document()
            for r in unique_results:
                doc.add_page_break()  # Start each use case on a new page
                
                # Parse the use case content
                paragraphs = r.split('\n\n')
                for para in paragraphs:
                    lines = para.strip().split('\n')
                    
                    # If this is just a single line
                    if len(lines) == 1:
                        line = lines[0].strip()
                        
                        # Title (in bold)
                        if line.startswith('**') and line.endswith('**') and not ':' in line:
                            doc.add_heading(line.strip('*'), level=1)
                        # Section headers
                        elif line.startswith('**') and ':**' in line:
                            p = doc.add_paragraph()
                            p.add_run(line).bold = True
                        # Regular paragraph
                        else:
                            doc.add_paragraph(line)
                    # Multi-line paragraph
                    else:
                        # Check if first line is a section header
                        first_line = lines[0].strip()
                        if first_line.startswith('**') and ':**' in first_line:
                            p = doc.add_paragraph()
                            p.add_run(first_line).bold = True
                            
                            # Add remaining lines as separate paragraphs
                            for line in lines[1:]:
                                if line.strip():
                                    if line.strip().startswith('-'):
                                        doc.add_paragraph(line.strip(), style='List Bullet')
                                    else:
                                        doc.add_paragraph(line.strip())
                        else:
                            # Just a regular multi-line paragraph
                            p = doc.add_paragraph()
                            for line in lines:
                                p.add_run(line + '\n')
            
            doc.save(summary_docx)
    except Exception as e:
        print(f"Error creating summary document: {e}")
    
    # Failed chunks
    with open(failed_txt, "w") as f:
        for idx in failed:
            f.write(f"Chunk {idx} failed\n")
    
    print(f"✅ Done: {len(unique_results)} use cases ({use_case_counter-1} total including low confidence), {len(failed)} failed chunks.")
    return len(unique_results), use_case_counter-1, len(failed)

def run_all():
    total_success = 0
    total_all = 0
    total_failed = 0
    
    print(f"Starting batch processing of {len(ALL_RPG_FILES)} files...")
    
    for fname in ALL_RPG_FILES:
        full_path = os.path.join(SOURCE_DIR, fname)
        program_match = re.search(r"AP(\d+)", fname.upper())
        
        if not program_match:
            print(f"Skipping {fname}: No AP### match.")
            continue
            
        if not os.path.exists(full_path):
            print(f"❌ File not found: {full_path}")
            continue
            
        program_num = program_match.group(1)
        print(f"\n{'=' * 60}\nProcessing file: {fname} (AP{program_num})\n{'=' * 60}")
        
        success, all_count, failed = process_chunks(full_path, program_num)
        total_success += success
        total_all += all_count
        total_failed += failed
    
    print(f"\n{'=' * 60}")
    print(f"BATCH PROCESSING COMPLETE")
    print(f"Total high quality use cases: {total_success}")
    print(f"Total use cases (including low confidence): {total_all}")
    print(f"Total failed chunks: {total_failed}")
    print(f"{'=' * 60}")

if __name__ == "__main__":
    # Create output directories if they don't exist
    os.makedirs(OUTPUT_BASE, exist_ok=True)
    os.makedirs("logs", exist_ok=True)
    
    if "--all" in sys.argv:
        run_all()
    elif len(sys.argv) >= 2:
        file_arg = sys.argv[1]
        program_match = re.search(r"AP(\d+)", file_arg.upper())
        if not program_match:
            print("❌ Invalid file name. Must include AP###.")
            sys.exit(1)
        process_chunks(file_arg, program_match.group(1))
    else:
        print("Usage: python analyze_rpg_usecases_final_batch.py [filename] or --all")