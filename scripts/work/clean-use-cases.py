import os
import docx
import shutil
from pathlib import Path

def clean_word_documents(source_folder, destination_folder, text_to_remove):
    """
    Opens Word documents from source folder, removes specified text, 
    and saves to destination folder
    """
    # Create destination folder if it doesn't exist
    if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)
        print(f"Created destination folder: {destination_folder}")
    
    # Get all docx files in the source folder
    source_path = Path(source_folder)
    docx_files = list(source_path.glob("*.docx")) + list(source_path.glob("*.doc"))
    
    if not docx_files:
        print("No Word documents found in the specified folder.")
        return
    
    print(f"Found {len(docx_files)} Word documents to process.")
    
    # Process each document
    for doc_path in docx_files:
        file_name = doc_path.name
        destination_path = os.path.join(destination_folder, file_name)
        
        print(f"Processing: {file_name}")
        
        try:
            # First copy the original file to the destination
            shutil.copy2(doc_path, destination_path)
            
            # Open the copied document
            doc = docx.Document(destination_path)
            
            # Flag to track if changes were made
            changes_made = False
            
            # Process each paragraph in the document
            for paragraph in doc.paragraphs:
                if text_to_remove in paragraph.text:
                    # Replace the text in the paragraph
                    new_text = paragraph.text.replace(text_to_remove, "")
                    paragraph.text = new_text
                    changes_made = True
            
            # Process text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            if text_to_remove in paragraph.text:
                                new_text = paragraph.text.replace(text_to_remove, "")
                                paragraph.text = new_text
                                changes_made = True
            
            # Save the document if changes were made
            if changes_made:
                doc.save(destination_path)
                print(f"  - Saved edited copy with changes to: {destination_path}")
            else:
                print(f"  - No instances of '{text_to_remove}' found in document")
                
        except Exception as e:
            print(f"Error processing file {file_name}: {str(e)}")
    
    print(f"Script completed. Edited files saved to: {destination_folder}")

if __name__ == "__main__":
    # Define paths and text to remove
    source_folder = r"C:\Users\Jeremy Callinan\Documents\Use-Cases\04072025-Submission\Unedited"
    destination_folder = r"C:\Users\Jeremy Callinan\Documents\Use-Cases\04072025-Submission\Edited"
    text_to_remove = "**"
    
    # Run the function
    clean_word_documents(source_folder, destination_folder, text_to_remove)