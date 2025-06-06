import os
import re
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document

PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TIMEOUT = 120
FUZZY_THRESHOLD = 0.94
DEBUG = True
INCLUDE_AP200 = False  # Toggle if needed

PRIMARY_FILE = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC\AP160.rpg36.txt"
CONTEXT_FILE = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC\AP200.rpg36.txt"
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"


def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f]


def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step)]


def build_prompt(template, chunk, context=None):
    context_block = f"\nContext code from AP200 (reference only):\n{context}\n" if context else ""
    return f"""{template}

You are analyzing **IBM RPG code from AP160** (Accounts Payable system). Your job is to extract only the business logic in a structured use case format. Focus on:
- AP vouchers, vendor validation, invoice rules, GL, payments, and 1099 logic.
- Skip implementation detail and low-level RPG structure.

DO NOT return summaries or commentary. Follow the format strictly.
{context_block}
[RPG CODE]
{chunk}
[END CODE]
"""


def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None


def normalize_headers(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    return text


def extract_use_case_number(text):
    match = re.search(r"\*\*Use Case ID\*\*:\s*UC-AP-\d+-([0-9]+)", text)
    return match.group(1) if match else "XXX"


def extract_title(text):
    title_match = re.search(r"^#\s*(.+)", text, re.MULTILINE)
    if title_match:
        return title_match.group(1).strip()
    desc_match = re.search(r"## Description\s+(.*?)\n", text, re.DOTALL)
    return desc_match.group(1).strip().split('.')[0] if desc_match else "Untitled"


def to_narrative(text):
    sections = ["Identification", "Description", "Pre-Condition", "Post-Condition",
                "Entities Used / Tables Used", "Program Steps", "Tests Needed"]
    out = ["Use Case Template\n"]
    for sec in sections:
        match = re.search(f"##+\s+{re.escape(sec)}(.*?)((?=## )|\Z)", text, re.DOTALL | re.IGNORECASE)
        if match:
            out.append(f"### {sec}\n")
            out.append(match.group(1).strip() + "\n")
    return '\n'.join(out)


def save_docx(text, path):
    doc = Document()
    for para in text.splitlines():
        if para.startswith("### "):
            doc.add_paragraph(para.replace("### ", "").strip()).bold = True
        else:
            doc.add_paragraph(para)
    doc.save(path)


def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique


def process_chunk(i, chunk, template, context_code, outdir, logdir, rawfile, docxdir):
    prompt = build_prompt(template, chunk, context_code)
    result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)

    if not result:
        with open(os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt"), "w") as f:
            f.write("❌ No output.")
        return None

    result = normalize_headers(result)
    with open(rawfile, "a", encoding="utf-8") as rf:
        rf.write(f"\n\n# Chunk {i+1}\n\n{result}\n\n{'='*60}\n")

    if "**Use Case ID**:" not in result or "## Description" not in result:
        with open(os.path.join(logdir, f"failed_chunk_{i+1:02d}.txt"), "w", encoding="utf-8") as f:
            f.write(result)
        return None

    uc_num = extract_use_case_number(result)
    title = extract_title(result)
    title_clean = re.sub(r'[^\w\d\- ]+', '', title).strip().replace(" ", "-")[:75]
    filename_base = f"UC-AP-160-{uc_num.zfill(3)}-{title_clean}"

    with open(os.path.join(outdir, f"{filename_base}.md"), "w", encoding="utf-8") as f:
        f.write(result + "\n")
    save_docx(result, os.path.join(docxdir, f"{filename_base}.docx"))
    print(f"✅ Chunk {i+1} saved as {filename_base}")
    return result


if __name__ == "__main__":
    lines = read_lines(PRIMARY_FILE)
    context = "\n".join(read_lines(CONTEXT_FILE)[:40]) if INCLUDE_AP200 else ""
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(lines)

    now = datetime.now()
    date_folder = os.path.join(OUTPUT_BASE, f"usecases-{now.strftime('%Y-%m-%d')}")
    run_folder = os.path.join(date_folder, f"mistral_debug_ap160_{now.strftime('%Y%m%d_%H%M%S')}")
    os.makedirs(run_folder, exist_ok=True)
    os.makedirs(os.path.join(run_folder, "word_docs"), exist_ok=True)
    os.makedirs(os.path.join("logs", f"logs_{now.strftime('%Y%m%d_%H%M%S')}"), exist_ok=True)

    rawfile = os.path.join(run_folder, "RAW_OLLAMA_OUTPUT.md")
    all_results = []
    for i, chunk in enumerate(chunks):
        res = process_chunk(i, chunk, template, context, run_folder,
                            os.path.join("logs", f"logs_{now.strftime('%Y%m%d_%H%M%S')}"),
                            rawfile,
                            os.path.join(run_folder, "word_docs"))
        if res:
            all_results.append(res)

    deduped = fuzzy_dedupe(all_results)
    with open(os.path.join(run_folder, "SUMMARY.md"), "w", encoding="utf-8") as f:
        for uc in deduped:
            f.write(to_narrative(uc) + "\n\n" + "=" * 60 + "\n\n")

    print(f"\n✅ {len(deduped)} strong use cases saved to: {run_folder}")
