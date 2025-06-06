import os
import re
import sys
import subprocess
from datetime import datetime
from docx import Document
from difflib import SequenceMatcher

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 120
MODULE = "AP"
MIN_CONFIDENCE_LENGTH = 400
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"

# === UTILITIES ===
def read_file(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f]

def save_as_docx(text, path):
    doc = Document()
    for line in text.splitlines():
        p = doc.add_paragraph()
        if "**" in line and ":" in line:
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(11)
                else:
                    p.add_run(part)
        else:
            p.add_run(line)
    doc.save(path)

def extract_title_and_id(text, program):
    title_match = re.search(r"#\s+(.*)", text)
    title = title_match.group(1).strip() if title_match else "Untitled"

    desc_match = re.search(r"## Description\s+(.*)", text, re.IGNORECASE)
    if not title or title.lower() == "untitled":
        title = desc_match.group(1).split(".")[0].strip() if desc_match else "Untitled"

    title_clean = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(" ", "-")[:75]

    id_match = re.search(r"UC-AP-[0-9]{3}-[0-9]{3}", text)
    use_case_id = id_match.group(0) if id_match else f"UC-AP-{program}-XXX"

    return use_case_id, title_clean or "Untitled"

def is_low_confidence(text):
    return len(text.strip()) < MIN_CONFIDENCE_LENGTH or "Untitled" in text or "Description" not in text

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def prompt_template(rpg_code, template, program):
    return f"""
You are analyzing **IBM RPG code** from program **AP{program}**, used in an **Accounts Payable** system.

Your task is to extract a business use case following this format:

{template}

🎯 Focus only on business logic like:
- Voucher creation, invoice processing, vendor lookup, check printing, GL validation
- Avoid low-level syntax or control flow

🧠 Be descriptive. Use meaningful field names. Return one structured use case.

[RPG CODE]
{rpg_code}
[END CODE]
"""

# === MAIN ===
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❌ Please specify the path to your RPG file (e.g., AP160.rpg36.txt)")
        sys.exit(1)

    rpg_file = sys.argv[1]
    program = re.findall(r"AP(\d+)", os.path.basename(rpg_file).upper())[0] if "AP" in rpg_file.upper() else "XXX"
    code = read_file(rpg_file)
    template = read_file(TEMPLATE_FILE)

    # Setup paths
    now = datetime.now()
    date_folder = f"usecases-{now.strftime('%Y-%m-%d')}"
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    base_path = os.path.join(OUTPUT_BASE, date_folder, f"{program}_{timestamp}")
    os.makedirs(base_path, exist_ok=True)
    docx_dir = os.path.join(base_path, "word_docs")
    os.makedirs(docx_dir, exist_ok=True)

    print(f"\n🚀 Running full superchunk analysis on {rpg_file}...\n")

    prompt = prompt_template(code, template, program)
    result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)

    if not result:
        print("❌ Model failed or timed out.")
        sys.exit(1)

    # Save outputs
    uc_id, title = extract_title_and_id(result, program)
    base_name = f"{uc_id}-{title}"
    base_name = re.sub(r"-+", "-", base_name)

    md_path = os.path.join(base_path, f"{base_name}.md")
    docx_path = os.path.join(docx_dir, f"{base_name}.docx")
    raw_log = os.path.join(base_path, "RAW_OUTPUT.md")
    review_doc = os.path.join(base_path, "SUMMARY_REVIEW.docx")

    with open(md_path, "w", encoding="utf-8") as f:
        f.write(result)
    with open(raw_log, "w", encoding="utf-8") as f:
        f.write(result)
    save_as_docx(result, docx_path)

    # Review Summary
    doc = Document()
    doc.add_heading("Use Case Summary Review", 0)
    doc.add_paragraph(f"Program: AP{program} — {uc_id}\n")
    if is_low_confidence(result):
        doc.add_paragraph("⚠️ Flagged as low confidence (short or vague)", style='Intense Quote')
    for line in result.splitlines():
        p = doc.add_paragraph()
        if line.startswith("##"):
            p.add_run(line.replace("##", "").strip()).bold = True
        else:
            p.add_run(line)
    doc.save(review_doc)

    print(f"✅ Done! Saved use case: {base_name}")
    print(f"📁 Output folder: {base_path}")
