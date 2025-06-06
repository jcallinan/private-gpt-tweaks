import os
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
import re
import sys
from docx import Document

PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
CHUNK_SIZE = 90
CHUNK_OVERLAP = 30
TIMEOUT = 120
DEBUG = True
FUZZY_THRESHOLD = 0.94
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"

def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i+size]) for i in range(0, len(lines), step)]

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def normalize_headers(text):
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    return text

def extract_title_and_id(text, program):
    match_id = re.search(r"(UC-[A-Z]+-\d+-\d+)", text)
    match_title = re.search(r"^#\s*(.*)", text, re.MULTILINE)
    uc_id = match_id.group(1).strip() if match_id else f"UC-AP-{program}-XXX"
    title = match_title.group(1).strip() if match_title else "Untitled"
    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-').strip('-')[:75]
    return uc_id, clean_title

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def format_narrative(text):
    sections = ["Identification", "Description", "Pre-Condition", "Post-Condition",
                "Entities Used / Tables Used", "Program Steps", "Tests Needed"]
    out = ["Use Case Template\n"]
    for sec in sections:
        match = re.search(f"##+\s+{re.escape(sec)}(.*?)((?=## )|\Z)", text, re.DOTALL | re.IGNORECASE)
        if match:
            out.append(f"### {sec}\n")
            out.append(match.group(1).strip() + "\n")
    return '\n'.join(out)

def save_as_docx(text, path):
    doc = Document()
    for line in text.splitlines():
        if line.startswith("###") or line.startswith("##"):
            doc.add_heading(line.replace("#", "").strip(), level=2)
        elif line.strip():
            doc.add_paragraph(line.strip())
        else:
            doc.add_paragraph("")
    doc.save(path)

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❌ Please provide an RPG source file as an argument.")
        sys.exit(1)

    source_file = sys.argv[1]
    if not os.path.exists(source_file):
        print(f"❌ File not found: {source_file}")
        sys.exit(1)

    program_match = re.search(r"AP(\d+)", os.path.basename(source_file).upper())
    program = program_match.group(1) if program_match else "XXX"

    lines = read_lines(source_file)
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(lines)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    date_dir = datetime.now().strftime("usecases-%Y-%m-%d")
    base = os.path.join(OUTPUT_BASE, date_dir, f"ap{program}_{timestamp}")
    os.makedirs(base, exist_ok=True)

    md_dir = os.path.join(base, "md")
    docx_dir = os.path.join(base, "docx")
    log_dir = os.path.join(base, "logs")
    os.makedirs(md_dir, exist_ok=True)
    os.makedirs(docx_dir, exist_ok=True)
    os.makedirs(log_dir, exist_ok=True)

    raw_path = os.path.join(base, "RAW.md")
    summary_path = os.path.join(base, "SUMMARY.md")
    failed_path = os.path.join(base, "FAILED_CHUNKS.txt")

    raw_log = open(raw_path, "w", encoding="utf-8")
    failed_chunks = []
    all_results = []

    for i, chunk in enumerate(chunks):
        print(f"🔍 Processing chunk {i+1}/{len(chunks)}...")
        prompt = f"{template}\n\nYou are analyzing IBM RPG code from AP{program}. Extract a structured use case. No commentary.\n\n[RPG CODE]\n{chunk}\n[END CODE]"
        result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)

        if not result:
            print(f"❌ No output for chunk {i+1}")
            failed_chunks.append(i+1)
            continue

        result = normalize_headers(result)
        raw_log.write(f"\n\n# Chunk {i+1}\n\n{result}\n\n{'='*60}\n")

        if "Use Case ID" in result and "## Identification" in result:
            uc_id, title = extract_title_and_id(result, program)
            fname = f"{uc_id}-{title}"
            md_path = os.path.join(md_dir, f"{fname}.md")
            docx_path = os.path.join(docx_dir, f"{fname}.docx")

            with open(md_path, "w", encoding="utf-8") as f:
                f.write(result)

            save_as_docx(result, docx_path)
            all_results.append(result)
            print(f"✅ Saved: {fname}")
        else:
            print(f"⚠️ Chunk {i+1} failed structure check.")
            with open(os.path.join(log_dir, f"failed_chunk_{i+1:02d}.txt"), "w", encoding="utf-8") as f:
                f.write(result)
            failed_chunks.append(i+1)

    raw_log.close()
    deduped = fuzzy_dedupe(all_results)

    with open(summary_path, "w", encoding="utf-8") as f:
        for uc in deduped:
            f.write(format_narrative(uc) + "\n\n" + "="*60 + "\n\n")

    with open(failed_path, "w", encoding="utf-8") as f:
        for idx in failed_chunks:
            f.write(f"Chunk {idx} failed\n")

    print(f"\n✅ Done! {len(deduped)} saved | {len(failed_chunks)} failed")
    print(f"📂 Output: {base}")
