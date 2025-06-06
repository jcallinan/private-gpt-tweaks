import os
import re
import sys
import subprocess
import time
from datetime import datetime
from difflib import SequenceMatcher
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 120
CHUNK_OVERLAP = 10
MIN_CONFIDENCE_LENGTH = 400
OUTPUT_BASE = "use_case_outputs"
MODULE_NAME = "AP"

TEMPLATE_FILE = "use_case_template.md"
AP200_PATH = "C:\\Temp\\IBM-GitHub-Submission-Unedited\\IBM-GitHub-Submission\\QSRC\\AP200.rpg36.txt"
INCLUDE_AP200 = False
DEBUG = True
FUZZY_THRESHOLD = 0.94


def read_lines(path):
    with open(path, "r", encoding="utf-8") as f:
        return [line.rstrip("\n") for line in f.readlines()]

def read_template(path):
    with open(path, "r", encoding="utf-8") as f:
        return f.read()

def fuzzy_match(a, b):
    return SequenceMatcher(None, a, b).ratio()

def is_low_confidence(text):
    # Relaxed rules to avoid overflagging
    required_sections = ["description", "process step", "validation", "pre-condition", "post-condition"]
    missing = sum(1 for s in required_sections if s not in text.lower())
    short = len(text.strip()) < MIN_CONFIDENCE_LENGTH
    return missing > 2 or short

def normalize_filename(title):
    safe = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(" ", "-").strip("-")
    return safe[:75]

def extract_use_case_id_and_title(text, program):
    id_match = re.search(r"UC-AP-[A-Z]*[-]?(AP)?\d+-\d+", text)
    title_match = re.search(r"#\s+(.*)", text)
    use_case_id = id_match.group(0).strip() if id_match else f"UC-AP-{program}-XXX"

    if title_match:
        title = title_match.group(1).strip()
    else:
        desc_match = re.search(r"## Description\s+(.*)", text, re.IGNORECASE)
        title = desc_match.group(1).split(".")[0].strip() if desc_match else "Untitled"

    clean_title = normalize_filename(title or "Untitled")
    return use_case_id, clean_title

def chunk_by_subroutine(lines):
    chunks = []
    buffer = []
    inside_sub = False
    for line in lines:
        if "BEGSR" in line:
            if buffer:
                chunks.append("\n".join(buffer))
                buffer = []
            inside_sub = True
        if inside_sub:
            buffer.append(line)
        if "ENDSR" in line:
            inside_sub = False
            chunks.append("\n".join(buffer))
            buffer = []
    if buffer:
        chunks.append("\n".join(buffer))
    return chunks

def build_prompt(template, code_chunk, context="", program="160"):
    context_block = f"\nContext code from AP200 (reference only):\n{context}\n" if context else ""
    return f"""{template}

You are analyzing **IBM RPG code from AP{program}** (Accounts Payable system). Your job is to identify the logic and generate a use case.

Start by summarizing the business logic. Then use the template format.

{context_block}

[RPG CODE]
{code_chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        return None

def save_as_docx(content, path):
    doc = Document()
    for line in content.splitlines():
        p = doc.add_paragraph()
        if "**" in line and ":" in line:
            parts = re.split(r"(\*\*.*?\*\*)", line)
            for part in parts:
                if part.startswith("**") and part.endswith("**"):
                    bold_run = p.add_run(part[2:-2])
                    bold_run.bold = True
                    bold_run.font.size = Pt(11)
                else:
                    p.add_run(part)
        else:
            p.add_run(line)
    doc.save(path)

# === MAIN ===
if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("❌ Please provide a primary RPG file path (e.g., AP160.rpg36.txt)")
        sys.exit(1)

    primary_file = sys.argv[1]
    program = re.findall(r"AP(\d+)", os.path.basename(primary_file).upper())[0] if "AP" in primary_file.upper() else "XXX"

    lines = read_lines(primary_file)
    ap200 = read_lines(AP200_PATH)[:40] if INCLUDE_AP200 else []
    template = read_template(TEMPLATE_FILE)
    chunks = chunk_by_subroutine(lines)

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    date_folder = os.path.join(OUTPUT_BASE, f"usecases-{datetime.now().strftime('%Y-%m-%d')}")
    output_dir = os.path.join(date_folder, f"{program.lower()}_{timestamp}")
    word_dir = os.path.join(output_dir, "word_docs")
    os.makedirs(word_dir, exist_ok=True)
    os.makedirs(output_dir, exist_ok=True)

    summary = []
    low_conf = []
    rawlog = os.path.join(output_dir, "RAW_OLLAMA_OUTPUT.md")
    failedlog = os.path.join(output_dir, "FAILED_CHUNKS.txt")

    with open(rawlog, "w", encoding="utf-8") as raw_out, open(failedlog, "w", encoding="utf-8") as failed_out:
        for i, chunk in enumerate(chunks):
            print(f"\n🔍 Chunk {i+1}/{len(chunks)}")
            prompt = build_prompt(template, chunk, "\n".join(ap200), program)
            result = run_ollama(PRIMARY_MODEL, prompt) or run_ollama(FALLBACK_MODEL, prompt)

            if not result:
                failed_out.write(f"❌ Chunk {i+1} timed out or failed.\n")
                continue

            raw_out.write(f"\n\n# Chunk {i+1}\n\n{result}\n{'='*60}\n")

            use_case_id, title = extract_use_case_id_and_title(result, program)
            base_name = f"{use_case_id}-{title}"

            md_path = os.path.join(output_dir, f"{base_name}.md")
            docx_path = os.path.join(word_dir, f"{base_name}.docx")

            try:
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)

                save_as_docx(result, docx_path)

                if is_low_confidence(result):
                    low_conf.append(result)
                    print(f"⚠️  Saved low-confidence use case: {base_name}")
                else:
                    summary.append(result)
                    print(f"✅  Saved use case: {base_name}")
            except Exception as e:
                failed_out.write(f"❌ Chunk {i+1} failed to write: {e}\n")

    def write_block(name, contents):
        if contents:
            md_path = os.path.join(output_dir, name + ".md")
            docx_path = os.path.join(output_dir, name + ".docx")
            with open(md_path, "w", encoding="utf-8") as f:
                for block in contents:
                    f.write(block + "\n\n" + "="*60 + "\n\n")
            save_as_docx("\n\n".join(contents), docx_path)

    write_block("SUMMARY", summary)
    write_block("LOW_CONFIDENCE_REVIEW", low_conf)

    print(f"\n✅ Done! {len(summary)} strong use cases, {len(low_conf)} low confidence, logs in {output_dir}")
