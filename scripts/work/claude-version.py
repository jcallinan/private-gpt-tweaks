import os
import re
import sys
import time
from datetime import datetime
from difflib import SequenceMatcher
import subprocess
from docx import Document

# === CONFIGURATION ===
PRIMARY_MODEL = "mistral:7b-instruct"
FALLBACK_MODEL = "codellama:13b-instruct-q4_K_M"
TIMEOUT = 60  # Reduced timeout for faster processing
CHUNK_SIZE = 90  # Original size
CHUNK_OVERLAP = 30  # Original overlap
FUZZY_THRESHOLD = 0.94
TEMPLATE_FILE = "use_case_template.md"
OUTPUT_BASE = "use_case_outputs"
INCLUDE_AP200 = False  # Simplified - no extra context
MIN_ACCEPTANCE_SCORE = 0.5  # Threshold for accepting use cases even if not perfect

# === ALL FILES TO PROCESS ===
ALL_RPG_FILES = [
    "AP160.rpg36.txt", "AP298.rpg36.txt", "AP105.rpg36.txt", "AP192.rpg36.txt", "AP290.rpg36.txt",
    "AP1099.rpg36.txt", "AP296.rpg36.txt", "AP991P.rpg36.txt", "AP315.rpg.txt", "AP316.rpg.txt",
    "AP760.rpg36.txt", "AP760P.rpg36.txt", "AP765.rpg36.txt", "AP765N.rpg36.txt", "AP780.rpg36.txt",
    "AP780P.rpg36.txt", "AP790.rpg36.txt", "AP910.rpgle.txt", "AP945.rpgle.txt", "AP945C.clp.txt",
    "AP3155.rpg.txt"
]
SOURCE_DIR = r"C:\Temp\IBM-GitHub-Submission-Unedited\IBM-GitHub-Submission\QSRC"

def read_lines(path):
    try:
        with open(path, "r", encoding="utf-8") as f:
            return [line.rstrip("\n") for line in f.readlines()]
    except UnicodeDecodeError:
        # Fallback to another encoding if utf-8 fails
        with open(path, "r", encoding="latin-1") as f:
            return [line.rstrip("\n") for line in f.readlines()]

def chunk_lines(lines, size=CHUNK_SIZE, overlap=CHUNK_OVERLAP):
    step = size - overlap
    return ['\n'.join(lines[i:i + size]) for i in range(0, len(lines), step)]

def build_prompt(template, chunk, context=None, program="XXX"):
    context_block = f"\nReference context from AP200 (not primary logic):\n{context}\n" if context else ""
    
    # Simplified prompt focused on getting output in the right format
    return f"""{template}
You are analyzing IBM RPG code from program AP{program}.
Extract meaningful business logic use cases related to: voucher handling, GL, invoice, payment processing, etc.

IMPORTANT REQUIREMENTS:
1. Use Case ID must follow format: UC-AP-{program}-XXX (where XXX is a sequential number starting at 001)
2. First line must be the title as '# <Clear descriptive title>'
3. Return ONLY structured output following the template exactly

{context_block}
[RPG CODE]
{chunk}
[END CODE]
"""

def run_ollama(model, prompt):
    try:
        result = subprocess.run(
            ["ollama", "run", model],
            input=prompt.encode(),
            stdout=subprocess.PIPE,
            stderr=subprocess.PIPE,
            timeout=TIMEOUT
        )
        return result.stdout.decode().strip()
    except subprocess.TimeoutExpired:
        print(f"Timeout using {model}")
        return None
    except Exception as e:
        print(f"Error running ollama: {e}")
        return None

def normalize_headers(text, program):
    # Basic header normalization
    replacements = {
        "## Input Validation": "## Input Type Validation Checks",
        "## Validation Rules": "## Input Type Validation Checks",
        "## Entities Used": "## Entities Used / Tables Used",
        "## Tables Used": "## Entities Used / Tables Used"
    }
    
    for wrong, correct in replacements.items():
        text = text.replace(wrong, correct)
    
    # Fix Use Case ID format and ensure it follows correct pattern
    text = re.sub(r"(Use Case ID\*\*: UC-AP)-(\d+)", rf"\1-{program}-\2", text)
    text = re.sub(r"(Use Case ID\*\*: UC-AP-{program}-)([^0-9])", r"\1001\2", text)
    
    # Make sure there's a title at the beginning if missing
    if not re.search(r"^#\s+", text, re.MULTILINE):
        text = "# AP" + program + " Use Case\n" + text
        
    return text

def calculate_acceptance_score(text):
    """Calculate a score for how complete/valid the use case is"""
    score = 0
    required_sections = [
        "Use Case ID", 
        "## Description", 
        "## Pre-Condition", 
        "## Post-Condition",
        "## Entities Used / Tables Used", 
        "## Program Steps", 
        "## Tests Needed"
    ]
    
    # Check sections
    for section in required_sections:
        if section in text:
            score += 0.1
            
            # Check content in sections except ID
            if section != "Use Case ID":
                pattern = f"{section}(.*?)(?=\n## |$)"
                match = re.search(pattern, text, re.DOTALL)
                if match and len(match.group(1).strip()) > 10:  # Some minimal content
                    score += 0.05
    
    # Check for title
    if re.search(r"^#\s+.{5,}", text, re.MULTILINE):
        score += 0.1
        
    return score

def is_valid_output(text):
    """Basic validation to check if this looks like a use case"""
    score = calculate_acceptance_score(text)
    return score >= MIN_ACCEPTANCE_SCORE

def extract_title(text, program):
    # Try to get the ID number
    match_id = re.search(r"Use Case ID.*?UC-AP.*?-(\d+)", text)
    id_part = match_id.group(1).zfill(3) if match_id else "001"
    
    # Try to get a title from the content
    title_match = re.search(r"(?i)^#\s+(.*?)$", text, re.MULTILINE)
    if title_match:
        title = title_match.group(1).strip()
    else:
        # If no title, look for description
        desc_match = re.search(r"##\s+Description\s+(.*?)(?=\n##|$)", text, re.DOTALL)
        title = desc_match.group(1).strip()[:50] if desc_match else "Use-Case"
    
    # Clean the title for use in filename
    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-')
    return f"UC-AP-{program}-{id_part}-{clean_title[:50]}"

def save_as_docx(content, path):
    try:
        doc = Document()
        
        # Extract title for first heading
        title_match = re.search(r"(?i)^#\s+(.*?)$", content, re.MULTILINE)
        if title_match:
            title = title_match.group(1).strip()
            doc.add_heading(title, level=1)
        
        # Process rest of content as paragraphs and headings
        lines = content.splitlines()
        for i, line in enumerate(lines):
            if i == 0 and line.startswith("# "):
                continue  # Skip title, already added above
                
            if line.startswith("## "):
                doc.add_heading(line[3:], level=2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], level=3)
            elif line.startswith("**") and "**:" in line:
                # Field labels like "**Use Case ID**:"
                parts = line.split(":", 1)
                if len(parts) == 2:
                    p = doc.add_paragraph()
                    p.add_run(parts[0].strip() + ":").bold = True
                    p.add_run(parts[1])
            else:
                doc.add_paragraph(line)
        
        doc.save(path)
        return True
    except Exception as e:
        print(f"Error saving Word document: {e}")
        return False

def fuzzy_dedupe(results):
    seen, unique = [], []
    for r in results:
        if not any(SequenceMatcher(None, r, s).ratio() > FUZZY_THRESHOLD for s in seen):
            seen.append(r)
            unique.append(r)
    return unique

def format_narrative(text):
    headers = ["Identification", "Description", "Pre-Condition", "Post-Condition",
               "Entities Used / Tables Used", "Program Steps", "Tests Needed"]
    result = ["Use Case Template\n"]
    for h in headers:
        m = re.search(f"##+\s+{re.escape(h)}(.*?)(?=\n## |$)", text, re.DOTALL)
        if m:
            result.append(f"### {h}\n{m.group(1).strip()}\n")
    return "\n".join(result)

def process_chunks(file_path, program):
    now = datetime.now()
    date_folder = os.path.join(OUTPUT_BASE, f"usecases-{now.strftime('%Y-%m-%d')}")
    timestamp = now.strftime("%Y%m%d_%H%M%S")
    output_dir = os.path.join(date_folder, f"{program}_{timestamp}")
    log_dir = os.path.join("logs", f"logs_{program}_{timestamp}")
    
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(log_dir, exist_ok=True)

    raw_out = os.path.join(output_dir, "RAW_OLLAMA_OUTPUT.md")
    summary_md = os.path.join(output_dir, "SUMMARY.md")
    summary_docx = os.path.join(output_dir, "SUMMARY.docx")
    failed_txt = os.path.join(output_dir, "FAILED_CHUNKS.txt")

    print(f"Processing {file_path} for program AP{program}...")
    
    all_lines = read_lines(file_path)
    template = open(TEMPLATE_FILE, "r", encoding="utf-8").read()
    chunks = chunk_lines(all_lines)
    
    print(f"Generated {len(chunks)} chunks to process")
    
    all_results = []
    failed = []
    raw_log = []
    use_case_counter = 1  # To ensure sequential numbering

    # Process each chunk
    for i, chunk in enumerate(chunks):
        print(f"Processing chunk {i+1}/{len(chunks)}")
        
        # Build and run prompt
        prompt = build_prompt(template, chunk, None, program)
        
        # Try primary model first
        print(f"Trying with {PRIMARY_MODEL}")
        result = run_ollama(PRIMARY_MODEL, prompt)
        
        # Fall back to secondary model if needed
        if not result:
            print(f"Falling back to {FALLBACK_MODEL}")
            result = run_ollama(FALLBACK_MODEL, prompt)
        
        # Process the result
        if result:
            # Normalize headers
            result = normalize_headers(result, program)
            raw_log.append(f"\n\n# Chunk {i + 1}\n{result}\n{'=' * 50}\n")
            
            # Check if it looks valid enough
            score = calculate_acceptance_score(result)
            
            if is_valid_output(result):
                # Make sure Use Case ID is in the right format with sequential numbering
                result = re.sub(r"(Use Case ID\*\*: UC-AP-{program}-)\d+", 
                              rf"\1{str(use_case_counter).zfill(3)}", result)
                
                # Generate filename with proper format: UC-AP-160-001-Create-Voucher.docx
                filename = extract_title(result, program)
                
                # If the filename doesn't have a good ID, fix it
                if not re.search(r"UC-AP-\d+-\d{3}-", filename):
                    filename = f"UC-AP-{program}-{str(use_case_counter).zfill(3)}-Use-Case"
                
                md_path = os.path.join(output_dir, f"{filename}.md")
                docx_path = os.path.join(output_dir, f"{filename}.docx")
                
                # Save as markdown
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)
                
                # Save as docx with better error handling
                if save_as_docx(result, docx_path):
                    print(f"✅ Created use case: {filename}")
                    all_results.append(result)
                    use_case_counter += 1
                else:
                    print(f"⚠️ Created use case markdown but Word doc failed: {filename}")
                    all_results.append(result)  # Still add it to results
                    use_case_counter += 1
            else:
                # For low confidence results, still try to extract title and save as regular use case
                # with a "Low-Confidence" prefix so they're easily identifiable
                base_filename = f"UC-AP-{program}-{str(use_case_counter).zfill(3)}-Low-Confidence"
                
                # If we can extract better title, use it
                title_match = re.search(r"(?i)^#\s+(.*?)$", result, re.MULTILINE)
                if title_match:
                    title = title_match.group(1).strip()
                    clean_title = re.sub(r'[^a-zA-Z0-9\- ]+', '', title).replace(' ', '-')
                    base_filename = f"UC-AP-{program}-{str(use_case_counter).zfill(3)}-{clean_title[:40]}"
                
                md_path = os.path.join(output_dir, f"{base_filename}.md")
                docx_path = os.path.join(output_dir, f"{base_filename}.docx")
                
                with open(md_path, "w", encoding="utf-8") as f:
                    f.write(result)
                    
                save_as_docx(result, docx_path)
                print(f"⚠️ Created low confidence use case: {base_filename}")
                use_case_counter += 1
        else:
            print(f"❌ Failed to process chunk {i+1}")
            failed.append(i + 1)
    
    # Write aggregated outputs
    print(f"Writing summary files...")
    
    # Raw output log
    with open(raw_out, "w", encoding="utf-8") as f:
        f.writelines(raw_log)
    
    # Deduped summary - always save this even if incomplete
    unique_results = fuzzy_dedupe(all_results)
    with open(summary_md, "w", encoding="utf-8") as f:
        for uc in unique_results:
            f.write(format_narrative(uc) + "\n\n" + "=" * 60 + "\n\n")
    
    # Summary as Word document - create with better error handling
    try:
        if unique_results:
            doc = Document()
            for r in unique_results:
                doc.add_paragraph("=" * 60)
                for line in format_narrative(r).splitlines():
                    if line.startswith("### "):
                        doc.add_heading(line[4:], level=2)
                    else:
                        doc.add_paragraph(line)
            doc.save(summary_docx)
    except Exception as e:
        print(f"Error creating summary document: {e}")
    
    # Failed chunks
    with open(failed_txt, "w") as f:
        for idx in failed:
            f.write(f"Chunk {idx} failed\n")
    
    print(f"✅ Done: {len(unique_results)} use cases ({use_case_counter-1} total including low confidence), {len(failed)} failed chunks.")
    return len(unique_results), use_case_counter-1, len(failed)

def run_all():
    total_success = 0
    total_all = 0
    total_failed = 0
    
    print(f"Starting batch processing of {len(ALL_RPG_FILES)} files...")
    
    for fname in ALL_RPG_FILES:
        full_path = os.path.join(SOURCE_DIR, fname)
        program_match = re.search(r"AP(\d+)", fname.upper())
        
        if not program_match:
            print(f"Skipping {fname}: No AP### match.")
            continue
            
        if not os.path.exists(full_path):
            print(f"❌ File not found: {full_path}")
            continue
            
        program_num = program_match.group(1)
        print(f"\n{'=' * 60}\nProcessing file: {fname} (AP{program_num})\n{'=' * 60}")
        
        success, all_count, failed = process_chunks(full_path, program_num)
        total_success += success
        total_all += all_count
        total_failed += failed
    
    print(f"\n{'=' * 60}")
    print(f"BATCH PROCESSING COMPLETE")
    print(f"Total high quality use cases: {total_success}")
    print(f"Total use cases (including low confidence): {total_all}")
    print(f"Total failed chunks: {total_failed}")
    print(f"{'=' * 60}")

if __name__ == "__main__":
    # Create output directories if they don't exist
    os.makedirs(OUTPUT_BASE, exist_ok=True)
    os.makedirs("logs", exist_ok=True)
    
    if "--all" in sys.argv:
        run_all()
    elif len(sys.argv) >= 2:
        file_arg = sys.argv[1]
        program_match = re.search(r"AP(\d+)", file_arg.upper())
        if not program_match:
            print("❌ Invalid file name. Must include AP###.")
            sys.exit(1)
        process_chunks(file_arg, program_match.group(1))
    else:
        print("Usage: python analyze_rpg_usecases_final_batch.py [filename] or --all")